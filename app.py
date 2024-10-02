from flask import Flask, render_template, request, redirect, url_for, session, send_file
import os
import sqlite3
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'supersecretkey'

# Folder paths
TICKETS_FOLDER = 'tickets/'
UPLOADS_FOLDER = 'uploads/'
os.makedirs(TICKETS_FOLDER, exist_ok=True)
os.makedirs(UPLOADS_FOLDER, exist_ok=True)

# Admin credentials
ADMIN_USERNAME = 'admin'
ADMIN_PASSWORD = 'C0nfid3nti@l'

# Initialize the SQLite database
def init_db():
    conn = sqlite3.connect('tickets.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tickets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            issue TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'Unchecked',  -- Default status
            filename TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

init_db()

def get_db_connection():
    conn = sqlite3.connect('tickets.db')
    conn.row_factory = sqlite3.Row
    return conn

# Function to create a DOCX file
def create_docx(name, issue, filename, ticket_filename):
    try:
        doc = Document()
        doc.add_heading('Ticket Information', level=1)
        doc.add_paragraph(f'Name: {name}')
        doc.add_paragraph(f'Issue: {issue}')

        if filename:
            doc.add_paragraph(f'Attached File: {filename}')
            img_path = os.path.join(UPLOADS_FOLDER, filename)
            if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                doc.add_picture(img_path, width=Inches(4.0))  # Adjust width as needed

        doc.save(os.path.join(TICKETS_FOLDER, ticket_filename))
        print(f"Ticket saved as DOCX: {ticket_filename}")  # Debug statement
    except Exception as e:
        print(f"Error creating DOCX file: {e}")  # Log the error
        raise

# Home route for ticket submission and admin login
@app.route('/')
def index():
    return render_template('index.html')

# Route for submitting tickets
@app.route('/submit_ticket', methods=['POST'])
def submit_ticket():
    try:
        name = request.form['name']
        issue = request.form['issue']
        file = request.files['file']

        sanitized_issue = issue[:10].replace(' ', '_').replace('/', '_')
        ticket_filename = f"{name}_{sanitized_issue}_ticket.docx"

        uploaded_file_path = None
        original_filename = None
        if file:
            original_filename = file.filename  # Store the original filename
            uploaded_file_path = os.path.join(UPLOADS_FOLDER, original_filename)
            file.save(uploaded_file_path)

        # Create the DOCX file
        create_docx(name, issue, original_filename if file else None, ticket_filename)

        # Store the ticket in the database with the original filename
        conn = get_db_connection()
        conn.execute('INSERT INTO tickets (name, issue, status, filename) VALUES (?, ?, ?, ?)',
                     (name, issue, 'Unchecked', ticket_filename))
        conn.commit()
        conn.close()

        return redirect(url_for('index'))
    except Exception as e:
        print(f"Error in submit_ticket: {e}")  # Log the error
        return "Internal Server Error", 500

# Admin login route
@app.route('/admin_login', methods=['POST'])
def admin_login():
    try:
        username = request.form['username']
        password = request.form['password']
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            session['admin_logged_in'] = True
            return redirect(url_for('admin_dashboard'))
        else:
            print("Invalid login attempt")  # Debug statement
            return redirect(url_for('index'))
    except Exception as e:
        print(f"Error in admin_login: {e}")  # Log the error
        return "Internal Server Error", 500

# Update ticket status route
@app.route('/update_status/<int:ticket_id>', methods=['POST'])
def update_status(ticket_id):
    try:
        new_status = request.form['status']
        conn = get_db_connection()
        conn.execute('UPDATE tickets SET status = ? WHERE id = ?', (new_status, ticket_id))
        conn.commit()
        conn.close()
        return redirect(url_for('admin_dashboard'))
    except Exception as e:
        print(f"Error updating status: {e}")  # Log the error
        return "Internal Server Error", 500

# Admin dashboard route
@app.route('/admin_dashboard')
def admin_dashboard():
    if not session.get('admin_logged_in'):
        return redirect(url_for('index'))

    tickets = []  # List to hold ticket information

    try:
        conn = get_db_connection()
        tickets_data = conn.execute('SELECT * FROM tickets').fetchall()
        for ticket in tickets_data:
            ticket_info = {
                'id': ticket['id'],  # Include ticket ID for status updates
                'name': ticket['name'],
                'issue': ticket['issue'],
                'status': ticket['status'],
                'file_link': ticket['filename']
            }
            tickets.append(ticket_info)
    except Exception as e:
        print(f"Error reading ticket files: {e}")  # Log the error

    return render_template('admin_dashboard.html', tickets=tickets)

# Route to delete a ticket
@app.route('/delete_ticket/<int:ticket_id>', methods=['POST'])
def delete_ticket(ticket_id):
    try:
        conn = get_db_connection()
        ticket = conn.execute('SELECT * FROM tickets WHERE id = ?', (ticket_id,)).fetchone()
        
        if ticket:
            ticket_filename = ticket['filename']
            ticket_file_path = os.path.join(TICKETS_FOLDER, ticket_filename)

            # Check if the ticket file exists and delete it
            if os.path.exists(ticket_file_path):
                os.remove(ticket_file_path)
                print(f"Deleted ticket file: {ticket_file_path}")  # Debug statement

            # Retrieve the original uploaded file name to delete the associated upload
            original_filename = ticket['filename']  # Use the correct field if you stored the original filename
            attachment_path = os.path.join(UPLOADS_FOLDER, original_filename)
            if os.path.exists(attachment_path):
                os.remove(attachment_path)
                print(f"Deleted associated upload: {attachment_path}")  # Debug statement

            # Delete the ticket record from the database
            conn.execute('DELETE FROM tickets WHERE id = ?', (ticket_id,))
            conn.commit()

        return redirect(url_for('admin_dashboard'))
    except Exception as e:
        print(f"Error deleting ticket: {e}")  # Log the error
        return "Internal Server Error", 500

# Route to clear all uploaded files
@app.route('/clear_uploads', methods=['POST'])
def clear_uploads():
    try:
        for filename in os.listdir(UPLOADS_FOLDER):
            file_path = os.path.join(UPLOADS_FOLDER, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
                print(f"Deleted upload file: {file_path}")  # Debug statement
        return redirect(url_for('admin_dashboard'))
    except Exception as e:
        print(f"Error clearing uploads: {e}")  # Log the error
        return "Internal Server Error", 500

# Route to clear all ticket documents
@app.route('/clear_tickets', methods=['POST'])
def clear_tickets():
    try:
        for filename in os.listdir(TICKETS_FOLDER):
            file_path = os.path.join(TICKETS_FOLDER, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)
                print(f"Deleted ticket file: {file_path}")  # Debug statement
        # Optionally, you might want to clear the database as well
        conn = get_db_connection()
        conn.execute('DELETE FROM tickets')
        conn.commit()
        conn.close()
        return redirect(url_for('admin_dashboard'))
    except Exception as e:
        print(f"Error clearing tickets: {e}")  # Log the error
        return "Internal Server Error", 500
        
# Route to delete a specific image from uploads
@app.route('/delete_image/<filename>', methods=['POST'])
def delete_image(filename):
    try:
        file_path = os.path.join(UPLOADS_FOLDER, filename)
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"Deleted image file: {file_path}")  # Debug statement
        return redirect(url_for('admin_dashboard'))
    except Exception as e:
        print(f"Error deleting image: {e}")  # Log the error
        return "Internal Server Error", 500

# Route to download a ticket DOCX
@app.route('/download_ticket/<filename>')
def download_ticket(filename):
    try:
        return send_file(os.path.join(TICKETS_FOLDER, filename), as_attachment=True)
    except Exception as e:
        print(f"Error downloading ticket: {e}")  # Log the error
        return "Internal Server Error", 500

# Admin logout route
@app.route('/logout')
def logout():
    session.pop('admin_logged_in', None)
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8888)  # Run the server on a specific IP
